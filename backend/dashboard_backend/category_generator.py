from datetime import datetime
import shutil
import pandas as pd
from pymongo import MongoClient
# from TendorAutomationTool.backend.dashboard_backend.utils1 import category_add_func
from .utils.celeryUtils import celeryUtils
from .utils.uploadFileFunc import uploadFileFunc
from .utils.utils import folder_creation, csvGeneratorFunc
from dashboard_backend.utils.gpt_utils import category_add_func
from azure.storage.blob import BlobServiceClient
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pymongo
# from .utils1 import *
from PyPDF2 import PdfWriter, PdfReader
from docx import Document
import openpyxl
import re
from authentication.email_utils import TenderEmailSender
from datetime import datetime
from pymongo.errors import PyMongoError
from .logger import audit_logger, exception_logger, transaction_logger
container_name = os.environ['container_name']
MONGODB_URI = os.environ['MONGODB_URI']
AZURE_STORAGE_CONNECTION_STRING = os.environ['AZURE_STORAGE_CONNECTION_STRING']


class CategoryGenerator:

    def __init__(self, data) -> None:

        self.tender_name = data.get("tender_name")
        self.tender_number = data.get("tender_number")
        self.division = data.get("division")
        self.published_date = data.get("published_date")
        self.file_type = data.get("file_type")
        self.file_path = data.get("file_path")
        self.uploaded_by = data.get("uploaded_by")
        self.folder_name = f"{self.tender_name}_{self.tender_number}_{self.published_date}"
        self.task_id = data.get("task_id")
        self.current_path = os.getcwd()

        # Initialize the Azure Blob Storage client
        self.blob_service_client = BlobServiceClient.from_connection_string(
            AZURE_STORAGE_CONNECTION_STRING
        )
        self.cur_datetime = datetime.now().astimezone()

    def save_merge_docx_to_mongodb(self, merge_docx_blob_path):
        try:
            client = MongoClient(MONGODB_URI)

            # Access a specific database
            db = client["Metadata"]
            collection = db["merge_docx_metadata"]

            # Define the query based on tender_number
            query = {"tender_number": self.tender_number}

            # Define the update data
            update_data = {
                "$set": {
                    "division": self.division,
                    "tender_name": self.tender_name,
                    "created_on": datetime.now(),
                    "created_by": self.uploaded_by,
                    "parameter_doc_blob_path": merge_docx_blob_path,
                }
            }

            # Use update_one with upsert=True to update existing or insert new
            result = collection.update_one(query, update_data, upsert=True)

            # Return the inserted or updated document's ID
            return result.upserted_id or result.modified_count

        except PyMongoError as mongo_err:
            # Handle PyMongo errors
            exception_logger.error("MongoDB Error:", mongo_err)
            return None
        except Exception as e:
            # Handle other exceptions
            exception_logger.error(e)
            return None
        finally:
            # Close the MongoDB connection in the 'finally' block
            client.close()

    def process_text(self, text):
        
        """
        Process the given text to split it into sections based on '**' and further into lines.
        Tables are detected and processed based on '|' symbol in lines.

        Args:
            text (str): The text to be processed.

        Returns:
            list: A list of lists, where each inner list represents a row of processed data.
        """
        
        try:
            # Split the text into sections based on '**'
            sections = text.split("**")
            processed_data = []

            for section in sections:
                print(section)
                # Split each section into lines
                lines = section.split("\n")
                table_processed = []
                for line in lines:
                    if '|' in line:
                        # Remove leading and trailing '|' and split the line into cells
                        cells = re.split(r'\s*\|\s*', line.strip('|'))
                        table_processed.append(cells)
                    else:
                        cells = re.split(r'\s*\|\s*', line.strip('\n'))
                        table_processed.append(cells)
                if table_processed:
                    processed_data.extend(table_processed)
                    # Add an empty row for spacing between tables
                    processed_data.append([""])

            # Remove the last empty row
            return processed_data[:-1]
        
        except Exception as e:
            exception_logger.error(f"Error processing text: {str(e)} ")

    # Function to write data to Excel
    def write_to_excel(self, data, filename='combined_tables.xlsx'):
        """
        Write the given data to an Excel file.

        Args:
            data (list): List of lists, where each inner list is a row to be written to Excel.
            filename (str): Name of the Excel file to be created.
        """
        try:
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Combined Tables"
            # print(len(data))
            for row_index, row in enumerate(data, start=1):
                for col_index, cell in enumerate(row, start=1):
                    ws.cell(row=row_index, column=col_index, value=cell)

            wb.save(filename)
        
        except Exception as e:
            exception_logger.error(f"Error processing text: {str(e)}")
            
            

    def create_doc_from_data(self, filename, data, c, x):
        """
        Create a Word document from the provided data.
        :param filename: The name of the file to save the document as.
        :param data: The string data to process into the document.
        :param c: A dictionary containing 'category'.
        :param x: A dictionary containing 'parameter'.
        """
        
        try:
            # Create a new document
            doc = Document()
            doc.add_heading(c["category"], level=1)
            doc.add_heading(x["parameter"], level=2)
            doc.add_paragraph(self.cur_datetime.isoformat())
            data = "\n"+data
            # Split the data into sections
            sections = data.split('\n\n')

            for section in sections:
                if '|' in section:
                    lines = section.split('\n')
                    headers = re.split(r'\s*\|\s*', lines[1].strip('| \n'))
                    num_cols = len(headers)

                    # Create a table in the document
                    table = doc.add_table(rows=1, cols=num_cols)

                    # Set the headers for the table
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header

                    # Fill the table with data
                    for line in lines[2:]:
                        if '|' in line:
                            row_cells = table.add_row().cells
                            values = re.split(r'\s*\|\s*', line.strip('| \n'))
                            for i in range(min(num_cols, len(values))):
                                row_cells[i].text = values[i]

                    # Apply a darker border to each cell
                    for row in table.rows:
                        for cell in row.cells:
                            for key in cell._element.xpath(".//w:tcBorders/*"):
                                key.set('w:sz', '24')  # border size (3 pt)
                                # border style (single line)
                                key.set('w:val', 'single')
                                # border color (black)
                                key.set('w:color', '000000')

                else:
                    paragraph = doc.add_paragraph(section)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        run.font.size = Pt(12)

            # Save the document
            doc.save(filename)
            
        except Exception as e:
            exception_logger.error(f"Error creating Word document: {str(e)}")

    def start_category_generator(self):
        audit_logger.info("Entered in category generator")

        client = MongoClient(MONGODB_URI)

        # Access a specific database
        db = client["Metadata"]

        # Access specific collections within the database
        collection = db["uploaded_tender_metadata"]
        tender_metadata = db["tender_metadata"]

        # Create an instance of the TenderEmailSender class
        email_sender = TenderEmailSender(self.uploaded_by, client, tender_metadata, self.tender_name, self.tender_number, collection)
        
        try:
            xlsx_blob_storage_string = ""
            xlsx_blob_path = ""
            # Initialize the MongoDB client

            celeryUtils().save_celery_task_metadata(self.tender_name, self.tender_number,self.task_id, task_status="category-processing")
            version = 0

            # Access specific collections within the database
            csv_collection = client.Metadata.csv_metadata

            output_collection = client.Metadata.output_metadata
            categories_collection = client.Metadata.categories
            categories_output_collection = client.Metadata.category_output_metadata
            version_collection = client.Metadata.versioning_category_metadata

            try:
                # Get the old category data from the 'output_metadata' collection
                old_category_data = output_collection.find(
                    {"tender_number": self.tender_number, "tender_name": self.tender_name})

                # Prepare the documents with additional fields and updated 'created_on'
                updated_documents = []
                for document in old_category_data:
                    updated_document = {
                        **document,
                        "uploaded_by": self.uploaded_by,  # Replace with the actual uploaded_by value
                        "created_on": self.cur_datetime
                    }
                    updated_documents.append(updated_document)

                # Insert the updated documents into the 'versioning_category_metadata' collection
                version_collection.insert_many(updated_documents)

                # Delete all documents in the 'output_metadata' collection that have the same 'tender_number' as the 'document'
                categories_collection.delete_many(
                    {"tender_number": self.tender_number, "tender_name": self.tender_name})

                if old_category_data.count() > 0:
                    if old_category_data[0]["version"]:
                        version = old_category_data[0]["version"]
                    else:
                        version = 0
                else:
                    version = 0

            except pymongo.errors.PyMongoError as mongo_err:
                # Handle PyMongo errors
                exception_logger.error("MongoDB Error:", mongo_err)

            except IndexError as idx_err:
                # Handle index errors (e.g., accessing old_category_data[0])
                exception_logger.error("Index Error:{}".format(str(idx_err)))

            except Exception as e:
                # Handle other exceptions
                exception_logger.error(e)

            # Fetch tender info and categories from the respective collections
            # logger.info(f"tender_number {self.tender_number}  tender_name {self.tender_name}")
            tender_info = list(csv_collection.find({
                "tender_number": str(self.tender_number),
                "tender_name": self.tender_name
            }))

            if len(tender_info) == 0:
                raise Exception("Not enough data for category generation")

            categories = list(categories_collection.find({}))
            df = pd.DataFrame(tender_info)

            final_document_detail = []
            file_document_detail_obj = {}

            # this holds the pages of pdfs that were used to create category
            # documents, we create a document for final category that has all
            # the pages that weren't used for any other category document.
            final_category_files = {}

            # Initialize the Azure Blob Storage client
            # blob_service_client = BlobServiceClient.from_connection_string(
            #     AZURE_STORAGE_CONNECTION_STRING
            # )

            csvGeneratorFunc().create_folder_if_not_exists(self.blob_service_client,f"{self.division}/{self.folder_name}/output_files/parameter_wise_doc/{self.cur_datetime}/",
            )

            csvGeneratorFunc().create_folder_if_not_exists(
                self.blob_service_client, f"{self.division}/{self.folder_name}/output_files/category_wise_doc/{self.cur_datetime}/")

            for idx, c in enumerate(categories):
                # audit_logger.info(f"=====> {categories}")
                if idx == 1:
                    break
                audit_logger.debug(f"remaining categories  ====>: {idx}/{len(categories)}")

                c_final_res = {
                    "category": c["category"],
                    "pdf_file": "", "parameters": []
                }

            # if c["category"] == "Scope of Work":
                file_document_detail_obj[c["category"]] = {}
                titles = []  # Dictionary to hold the pages and pdf names from which we need to slice pages
                filtered_rows_other_formats = []

                for idx, x in enumerate(c["sub_category"]):
                    # Summarize output using GPT
                    summarized = category_add_func().summarized_gpt(
                        df, x["summarization_prompt"], x["gpt_prompt"])

                    # Extract the titles that are referred to for GPT output, we'll later extract all the pages under these titles
                    for z in summarized["filtered"]:
                        titles.append(z["title"])
                    # Create a document for this parameter with the data GPT returns
                    docx_file_name = (
                        x["parameter"] + "-" + self.cur_datetime.isoformat() +".docx")
                    
                    folder_creation(os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{self.tender_name}_{self.tender_number}", "docx_file"))
                    parameter_doc = (os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{self.tender_name}_{self.tender_number}", "docx_file", docx_file_name))

                    if "|" in summarized["gpt_output_gpt4"]:
                       
                        xlsx_file_name = parameter_doc.replace(
                            ".docx", ".xlsx")
                        # write xlsx file
                        processed_data = self.process_text(
                            summarized["gpt_output_gpt4"])
                        self.write_to_excel(processed_data, filename=xlsx_file_name)
                        # write doc file
                        self.create_doc_from_data(parameter_doc, summarized["gpt_output_gpt4"], c, x)
                        # Generate the blob storage path for the XLSX file
                        xlsx_blob_storage_string = f"{self.division}/{self.folder_name}/output_files/parameter_wise_doc/{self.cur_datetime}/{xlsx_file_name}"
                        xlsx_blob_path = category_add_func().saving_docx_to_blob(self.blob_service_client, 
                                                xlsx_blob_storage_string, xlsx_file_name)
                    else:

                        document = Document()
                        document.add_heading(c["category"], level=1)
                        document.add_heading(x["parameter"], level=2)
                        document.add_paragraph(self.cur_datetime.isoformat())
                        document.add_paragraph(summarized["gpt_output_gpt4"])
                        document.save(parameter_doc)

                    # Append the generated docx file path to the list
                    # individual_docx_files.append(parameter_doc)
                    parameter_blob_storage_string = f"{self.division}/{self.folder_name}/output_files/parameter_wise_doc/{self.cur_datetime}/{docx_file_name}"
                    docx_blob_path = category_add_func().saving_docx_to_blob(
                        self.blob_service_client, parameter_blob_storage_string, parameter_doc
                    )
                    # Save the output data to MongoDB
                    output_data = output_collection.insert_one(
                        {
                            "division": self.division,
                            "tender_number": self.tender_number,
                            "tender_name": tender_info[0]["tender_name"],
                            "gpt_output": summarized["gpt_output"],
                            "category": c["category"],
                            "parameter": x["parameter"],
                            "created_on": self.cur_datetime,
                            "created_by": self.uploaded_by,
                            "updated_on": "",
                            "updated_by": "",
                            "parameter_doc_blob_path": parameter_blob_storage_string,
                            "parameter_xlsx_blob_path": xlsx_blob_storage_string,
                            "version": 0,
                        }
                    )
                    # Append relevant data to the category response
                    c_final_res["parameters"].append(
                        {
                            "division": self.division,
                            "tender_number": self.tender_number,
                            "tender_name": tender_info[0]["tender_name"],
                            "gpt_output": summarized["gpt_output"],
                            "category": c["category"],
                            "parameter": x["parameter"],
                            "created_on": self.cur_datetime,
                            "parameter_doc_blob_path": parameter_blob_storage_string,
                            "parameter_xlsx_blob_path": xlsx_blob_storage_string
                        }
                    )
                    # Append relevant data to the file document detail object
                    file_document_detail_obj[c["category"]][x["parameter"]] = {
                        "id": str(output_data.inserted_id),
                        "division": self.division,
                        "tender_number": self.tender_number,
                        "tender_name": tender_info[0]["tender_name"],
                        "gpt_output": summarized["gpt_output"],
                        "category": c["category"],
                        "parameter": x["parameter"],
                        "created_on": self.cur_datetime,
                        "parameter_doc_blob_path": docx_blob_path,
                        "parameter_xlsx_blob_path": xlsx_blob_path

                    }

                # Here we create a pdf for the entire category
                pdf_file_name = c["category"] + "-" + \
                    self.cur_datetime.isoformat() + ".pdf"
                # Extract all the titles in all the pdfs that have been used
                titles = set(titles)
                pages = {}
                # Now we need to extract all the page numbers under those titles
                for t in tender_info:
                    # if t['title'] in titles:
                    if pages.get(t['tender_file_path']):
                        pages[t["tender_file_path"]].append(int(t["page"]))
                    else:

                        pages[t["tender_file_path"]] = [int(t["page"]), ]
                # Get only unique page numbers
                for file in list(pages.keys()):
                    pages[file] = list(set(pages[file]))
                    pages[file].sort()

                # Output file for this category
                output = PdfWriter()
                # Add all the pages to output
                for file in list(pages.keys()):

                    local_file = csvGeneratorFunc().download_file(
                        file, file.split(
                            "/")[-1], self.tender_name, self.tender_number
                    )
                    if local_file.endswith(('.docx', '.xlsx', 'doc', '.xls', 'DOCX', 'XLSX', 'XLS', 'DOC')):
                        local_file = os.path.splitext(local_file)[0] + '.pdf'

                    file_reader = PdfReader(open(local_file, "rb"))
                    for p in pages[file]:
                        output.add_page(file_reader.pages[p])
                        if final_category_files.get(file):
                            final_category_files[file].append(int(p))
                        else:
                            final_category_files[file] = [int(p), ]

                folder_creation(os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{self.tender_name}_{self.tender_number}", "docx_file"))
                category_doc = os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{self.tender_name}_{self.tender_number}", "docx_file", c["category"] + ".pdf"
                )
                outputStream = open(category_doc, "wb")
                output.write(outputStream)
                outputStream.close()
                category_blob_storage_string = (
                    f"{self.division}/{self.folder_name}/output_files/category_wise_doc/"
                )
                csvGeneratorFunc().create_folder_if_not_exists(
                    self.blob_service_client, category_blob_storage_string
                )
                csvGeneratorFunc().create_folder_if_not_exists(
                    self.blob_service_client, f"{category_blob_storage_string}/{self.cur_datetime}/"
                )
                # Generate the PDF file path
                category_wise_pdf_path = f"{self.division}/{self.folder_name}/output_files/category_wise_doc/{self.cur_datetime}/{pdf_file_name}"
                # Save the category output data to MongoDB
                cat_output = categories_output_collection.insert_one(
                    {
                        "tender_name": self.tender_name,
                        "category_name": c["category"],
                        "tender_number": self.tender_number,
                        "pdf_file_path": category_wise_pdf_path,
                        'created_on': self.cur_datetime,
                        'created_by': self.uploaded_by,
                        "updated_on": "",
                        "updated_by": "",
                    }
                )
                # Update the file_document_detail_obj with category output idF
                file_document_detail_obj[c["category"]]["category_output_id"] = str(
                    cat_output.inserted_id
                )
                pdf_blob_path = category_add_func().saving_docx_to_blob(
                    self.blob_service_client, category_wise_pdf_path, category_doc
                )
                # Update the category response with the pdf_blob_path
                c_final_res["pdf_file"] = pdf_blob_path
                # Append the final result for this category to the list
                final_document_detail.append(c_final_res)

            # final category document creation
            output = PdfWriter()

            # these are the pages that have been used in other categories, we
            # don't want them in final category.
            pages = final_category_files

            # Add all the pages to output
            for file in list(pages.keys()):
                local_file = csvGeneratorFunc().download_file(
                    file, file.split(
                        "/")[-1], self.tender_name, self.tender_number
                )
                # audit_logger.info(f"local_file ===> {local_file}")

                if local_file.endswith(('.docx', '.xlsx', 'doc', '.xls', 'DOCX', 'XLSX', 'XLS', 'DOC')):
                    local_file = os.path.splitext(local_file)[0] + '.pdf'
                # audit_logger.info(f"local_file ===> {local_file}")

                file_reader = PdfReader(open(local_file, "rb"))
                pages[file] = set(pages[file])
                for p in range(0, len(file_reader.pages)):
                    # if page is not used in any other category wise doc creation.
                    if p not in pages[file]:
                        output.add_page(file_reader.pages[p])

            pdf_file_name = "final_category_doc" + "-" + self.cur_datetime.isoformat() + \
                ".pdf"
            category_doc = os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{self.tender_name}_{self.tender_number}", "docx_file", "final_category_doc.pdf")
            outputStream = open(category_doc, "wb")
            output.write(outputStream)
            outputStream.close()

            csvGeneratorFunc().create_folder_if_not_exists(
                self.blob_service_client, f"{self.division}/{self.folder_name}/output_files/category_wise_doc/{self.cur_datetime}/"
            )
            # Generate the PDF file path
            category_wise_pdf_path = f"{self.division}/{self.folder_name}/output_files/category_wise_doc/{self.cur_datetime}/{pdf_file_name}"
            # Save the category output data to MongoDB
            cat_output = categories_output_collection.insert_one(
                {
                    "tender_name": self.tender_name,
                    "category_name": "final_category_doc",
                    "tender_number": self.tender_number,
                    "pdf_file_path": category_wise_pdf_path,
                    'created_on': self.cur_datetime,
                    'created_by': self.uploaded_by,
                    "updated_on": "",
                    "updated_by": "",
                    # "merger_docx_file_id":merge_doc_file_id

                }
            )

            # Update the file_document_detail_obj with category output id
            file_document_detail_obj["final_category_doc"] = {}
            file_document_detail_obj['final_category_doc']["category_output_id"] = str(
                cat_output.inserted_id
            )
            pdf_blob_path = category_add_func().saving_docx_to_blob(
                self.blob_service_client, category_wise_pdf_path, category_doc
            )
            # Update the category response with the pdf_blob_path
            c_final_res["pdf_file"] = pdf_blob_path
            # Append the final result for this category to the list
            final_document_detail.append(c_final_res)

            category_list = [
                "LDDate", "Project Timeline", "Project Introduction", "LD Clause", "contact", "SOW Battery Limits", "SOW Commissioning", "SOW Design", "SOW Exclusion", "SOW Installation", "SOW O&M", "SOW Operational Support", "SOW PGTest", "SOW Supply", "SOW TrainingReq", "SOW TrialRunSpecs", "SOW TerminalPoints", "SOW SpecialRequirements", "SOW DesignCodes",
                "Input parameters", "source of water", "output parameter", "Critical Parameters", "Treatment_Scheme",
                "Location_and_Layout", "plant capacity", "operational hours", "position_moc_motor", "hrscc_specifications",
                "Chemical Dosing Tank", "Quantity of Dosing Pumps", "Specifications of Dosing Pumps", "Dosing Quantity", "pH Correction or Boosting System", "Auto Handling System",
                "bulk caustic handling system", "bulk hydrochloric acid handling system", "bulk sulphuric acid handling system", "other chemical bulk handling system", "unloading pumps",
                "MGF or PSF or DMF", "ACF or LLACF", "IRF", "SSF", "Mechanical Bought out", "Electrical and Instrumentation", "Piping and Valves Makes", "DM Plant", "Softener", "UF unit", "RO unit",
                "Storage Tanks specifications", "Other Tanks", "Centrifugal Pumps", "Sludge Transfer Pumps", "Line Description", "Pipe Specifications", "Support and Fitting",
                "Butterfly Valve", "Diaphragm Valve", "Ball Valve", "Globe Valve", "Gate Valve", "Check Valve", "Other Valves",
                "mcc_panel_specifications", "motor specifications", "Power Cables", "Instrumentation and Control Cables",
                "cable_tray_specifications", "LPBS_specifications", "earthing_specifications", "Quality Control Procedure", "Quality Assurance Plan", "Documentation", "Test Certificate", "Inspection Procedure",
                "Transmitter", "Quality measuring instruments", "Switch", "Gauge", "Solenoid valve box", "Other Instruments",
                "Welding Specifications", "Painting Specifications", "Warranty", "Drawing and Documents to be submitted along with the offer", "Drawing and Documents to be submitted after award of contract", "Material Packaging and Shipment", "Spares"
            ]

            merged_docx_path, merge_docx_file_name = category_add_func().merge_documents(
                category_list, self.tender_name, self.cur_datetime, self.tender_number)
            merge_docx_blob_storage_string = f"{self.division}/{self.folder_name}/output_files/general_docx/{merge_docx_file_name}"

            merge_docx_blob_path = category_add_func().saving_docx_to_blob(self.blob_service_client, merge_docx_blob_storage_string, merged_docx_path)

            merge_doc_file_id = self.save_merge_docx_to_mongodb(
                merge_docx_blob_storage_string)

            # Update tender metadata with the status
            uploadFileFunc().update_tender_metadata(
                self.tender_name, self.tender_number, self.division, self.uploaded_by, tender_status="Succeeded"
            )

            category_add_func().update_file_processing_status(
                self.tender_name, self.tender_number, self.uploaded_by, file_processing_status='Succeeded'
            )

            # To send a completion email
            #uery the MongoDB collection for a specific tender
            tender_status_doc = tender_metadata.find_one({"tender_number": self.tender_number, "tender_name": self.tender_name})

            # Check if a document was found and then check the tender status
            if tender_status_doc and tender_status_doc.get("status") == "Succeeded":
                # Call the cleanup method if the tender status is 'Succeeded'
                split_folder=os.path.join(f"/mnt/supporting_folders/{self.tender_name}_{self.tender_number}/splitted_pdf")
                sample_doc =os.path.join(f"/mnt/supporting_folders/{self.tender_name}_{self.tender_number}/sample_doc")
                sample_dataset =os.path.join(f"/mnt/supporting_folders/{self.tender_name}_{self.tender_number}/sample_dataset")
                docx_file = os.path.join(f"/mnt/supporting_folders/{self.tender_name}_{self.tender_number}/docx_file")
                general_docx_file = os.path.join(f"/mnt/supporting_folders/{self.tender_name}_{self.tender_number}/general_docx")

        
                shutil.rmtree(split_folder)
                shutil.rmtree(sample_doc)
                shutil.rmtree(sample_dataset)
                shutil.rmtree(docx_file)
                shutil.rmtree(general_docx_file)


            email_sender.send_completion_email()


            # Return the response with the final document details
            return {
                "message": "category wise document created",
                "status": "Succeeded",
                "file_document_detail": file_document_detail_obj,
                "final_document_detail": final_document_detail,
                "merger_docx_file_id": merge_doc_file_id
            }

            # Use the exception_logger to log any exceptions
        except Exception as e:
            exception_logger.exception(
                f"An error occurred during category wise document creation  {str(e)}")

            # Handle any other uncaught exceptions
            uploadFileFunc().update_tender_metadata(
                self.tender_name, self.tender_number, self.division, self.uploaded_by, tender_status='Failed'
            )
            category_add_func().update_file_processing_status(
                self.tender_name, self.tender_number, self.uploaded_by, file_processing_status="Failed")

            # To send a failure email
            email_msg = "An issue has been encountered during the category generation process. Please reach out to the development team for advanced technical assistance"
            if "no embeddings were generated" in str(e):
                email_sender.send_failure_email(str(e))
            elif "Not enough data for category generation" in str(e):
                email_sender.send_failure_email("Not enough data for category generation")
            else:
                email_sender.send_failure_email(email_msg)


            # Handle any exceptions that occur during the process and return an error response with status 500
            return {
                "message": "An error occurred during category wise document creation",
                "status": 500,
                "error": str(e),
            }
