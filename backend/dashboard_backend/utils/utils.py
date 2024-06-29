from PyPDF2 import PdfFileReader, PdfFileWriter, PdfWriter, PdfReader
import PyPDF2

from .adobeKeyUtils import adobeKeyUtils
from .PagesCount import PagesCount
import os
import csv
import openai
import time
import tiktoken
import io
import pandas as pd
from dashboard_backend.utils.gpt_utils import category_add_func
from dashboard_backend.utils.csv_utils.databasecreator import databasecreator
from dashboard_backend.utils.uploadFileFunc import uploadFileFunc
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader, PdfWriter
from pymongo import MongoClient
from pymongo.errors import PyMongoError
from PIL import Image
from django.core.files.storage import FileSystemStorage
from azure.storage.blob import BlobServiceClient
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.units import inch
from docx import Document
import fitz
import openai.error
from ..logger import *
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import mimetypes

AZURE_STORAGE_CONNECTION_STRING = os.environ['AZURE_STORAGE_CONNECTION_STRING']
MONGODB_URI = os.environ['MONGODB_URI']
container_name = os.environ['container_name']
openai_key = os.environ['openai_key']

ADOBE_EMAIL = os.environ['ADOBE_EMAIL']
DEVELOPER_EMAIL_ADOBE_KEY_EXPIRES = os.environ['DEVELOPER_EMAIL_ADOBE_KEY_EXPIRES']
openai.api_key = openai_key


def folder_creation(local_file_path):
    if not os.path.exists(local_file_path):
        os.makedirs(local_file_path)


def create_pdf(pdf_file, large_string):
    doc = SimpleDocTemplate(pdf_file, pagesize=letter)

    # Create a stylesheet
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    # Split the large string into paragraphs
    paragraphs = large_string.split('\n\n')

    # Build the content for each page
    content = []
    for paragraph in paragraphs:
        content.append(Paragraph(paragraph, normal_style))
        content.append(Spacer(1, 12))  # Add some space between paragraphs

    # Build the PDF document with multiple pages
    doc.build(content)


class document_func():

    def __init__(self):
        self.max_tokens = 1500
        self.model_engine = "text-davinci-003"
        self.encoding = tiktoken.get_encoding("cl100k_base")
        self.current_path = os.getcwd()

    def getText(self, doc):
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return fullText

    def create_chunk(self, input_string):
        """
        Creates chunks of strings based on the maximum number of tokens, splitting at the sentence level.
        """
        chunks = []  # List to store the generated chunks of text
        current_chunk = ""  # Current chunk being built
        current_token_count = 0  # Current count of tokens in the chunk

        # Split the input_string into sentences
        sentences = input_string.split('. ')

        for sentence in sentences:
            # Get the number of tokens in the current sentence
            sentence_tokens = self.encoding.encode(
                sentence, disallowed_special=())
            sentence_token_count = len(sentence_tokens)

            # Check if adding the current sentence exceeds the maximum token count
            if current_token_count + sentence_token_count > self.max_tokens:
                chunks.append({"text": current_chunk.strip()})
                # Append the current chunk as a dictionary
                current_chunk = ""  # Reset the current chunk
                current_token_count = 0  # Reset the token count

            current_chunk += sentence + ". "  # Add the sentence to the current chunk
            current_token_count += sentence_token_count  # Update the token count

            # Check if the current chunk exceeds the model's maximum context length
            if current_token_count > self.max_tokens:
                chunks.append({"text": current_chunk.strip()})
                # Append the current chunk as a dictionary
                current_chunk = ""  # Reset the current chunk
                current_token_count = 0  # Reset the token count

        # Append any remaining chunk
        if current_chunk:
            chunks.append({"text": current_chunk.strip()})

        return chunks

    
    
    def remove_a3_5_from_pdf(self,file_path,tender_name,file_name,tender_number):
        # audit_logger.info(f"====> {file_path}")
        audit_logger.info(f"====> {file_name}")
        folder_creation(os.path.join(
                "/mnt/supporting_folders",f"{tender_name}_{tender_number}", "uploaded_pdfs_converted"))
        output_file_path = os.path.join('/mnt/supporting_folders', f"{tender_name}_{tender_number}", 'uploaded_pdfs_converted', file_name+".pdf")
        pdf = PdfReader(open(file_path, 'rb'))
        page_sizes = []
        writer = PdfWriter()
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            size = (page.mediabox.width, page.mediabox.height)
            if not (size [0]>1000 and size[1]>1000):
                writer.add_page(page)
        with open(output_file_path, 'wb') as output_file:
            writer.write(output_file)
        
        
        return output_file_path 

class Ingest():

    def __init__(self):
        self.EMBEDDING_MODEL = "text-embedding-ada-002"
        openai.api_key = openai_key
        self.current_path = os.getcwd()

    def get_embedding(self, text, model="text-embedding-ada-002"):
        '''
        Creates openai embedding for the csv generated from pdf.
        '''
        max_retries = 5

        for attempt in range(max_retries):
            try:
                text = text.replace("\n", " ").replace("start_para_", "").replace(
                    "Title", "").replace("header", "").replace("subheader", "")
                res = openai.Embedding.create(input=[text], model=model)[
                    'data'][0]['embedding']
                return res
            except Exception as e:
                exception_logger.error(
                    "GPT error, waiting 90 seconds before retrying", flush=True)

                if attempt >= max_retries - 1:  # Check if it's the last attempt
                    audit_logger.info("OpenAI key limit exceeded")
                    exception_logger.exception(
                        "OpenAI API down, please contact technical team")
                    raise e  # Raise the exception if all retries are exhausted

                time.sleep(60)  # Wait for 90 seconds before retrying

        # If the loop completes without a successful return, raise an exception.
        raise Exception("Failed to create embedding after several retries.")

        # return {"error":"error in getting embeddings"}
        # return ({"error":"OpenAi Key limit exits "})

    def compute_doc_embeddings(self, df: pd.DataFrame):
        return {idx: self.get_embedding(str(r['text'])) for idx, r in df.iterrows()}

    def load_embeddings(self, df):
        max_dim = max([int(c) for c in df.columns if c != "count"])
        return {
            (r.count): [r[str(i)] for i in range(max_dim + 1)] for _, r in df.iterrows()
        }

    def process(self, data, mydir, type='pdf'):
        process_retry_count = 0
        max_process_retries = 5
        embedding_retry_count = 0
        max_embedding_retries = 10

        while process_retry_count < max_process_retries:
            try:
                document_embeddings = self.compute_doc_embeddings(data)
                final_df = pd.DataFrame(list(document_embeddings.items()), columns=[
                                        'Index', 'embedding'])
                final_df['Index'] = final_df['Index'].astype(int)

                # Add columns based on type
                if type == 'pdf':
                    final_df['title'] = data['title']
                    final_df['sub_title'] = data['sub_title']
                    final_df['text'] = data['text']
                    final_df['page'] = data['page']

                elif type in ['doc', 'excel']:
                    final_df['title'] = ""
                    final_df['sub_title'] = ""
                    final_df['text'] = data['text']
                    final_df['page'] = 0

                final_df['filename'] = (
                    mydir.rsplit("/")[-1]).rsplit(".csv")[0]

                # Check for required columns
                if len(final_df.columns) != 7 or 'embedding' not in final_df.columns:
                    audit_logger.warning(f"Incorrect number of columns or 'embedding' column missing. Columns: {final_df.columns}")
                    if 'embedding' not in final_df.columns and embedding_retry_count < max_embedding_retries:
                        embedding_retry_count += 1
                        time.sleep(60)  # Short pause before retry
                        continue
                    else:
                        raise e
                        raise ValueError( "Dataframe format issue after retries.")

                # Save to CSV
                audit_logger.debug(f"No of Columns generated in dataframe : {final_df.columns}")
                
                with open(mydir, 'a', encoding='utf-8') as f:
                    final_df.to_csv(f, header=f.tell() == 0, index=False)
                audit_logger.debug("Dataframe processed successfully.")
                return final_df

            except Exception as e:
                exception_logger.error(f"Error during processing: {str(e)}")
                process_retry_count += 1
                if process_retry_count >= max_process_retries:
                    exception_logger.exception( "Maximum retries reached for processing.")
                    raise e
        raise e
        raise Exception("Failed to process after all retries.")


class split_tender():
    def __init__(self):
        self.current_path = os.getcwd()

    def saving_file_to_folder(self, file, tender_name, tender_number, tender=False):
        folder = os.path.join(
            os.getcwd(), '/mnt/supporting_folders', f"{tender_name}_{tender_number}", 'amendments')
        filenametimeframe = databasecreator().generate_file_name(flag='pdf')
        fs = FileSystemStorage(location=folder)
        filename = fs.save(filenametimeframe, file)
        file_url = fs.url(filename)
        file_path = os.path.join(folder, filenametimeframe)
        return file_path

    def mod_data_fetch_from_amendment(self, file_path, tender_name, tender_number):
        obj = PdfFileReader(file_path)
        pgno = obj.getNumPages()
        output = PdfFileWriter()

        mod_list = []

        for i in range(0, pgno):
            PgOb = obj.getPage(i)
            Text = PgOb.extractText()
            if "Modifications/Additions/Deletions" in Text and i < 20:
                mod_list.append(i)
            else:
                pass

        # print(mod_list)

        for i in mod_list:
            PgOb = obj.getPage(i)
            output.addPage(PgOb)

        extracted_file_path = os.path.join(
            self.current_path, "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "splitted_amed", "extracted_amendments.pdf")
        folder_creation(extracted_file_path)
        with open(extracted_file_path, 'wb') as f:
            output.write(f)
        print(extracted_file_path)
        return extracted_file_path


class csvGeneratorFunc():

    def __init__(self):
        self.current_path = os.getcwd()
        self.A3_A5_0_thresh = 1000
        self.A3_A5_1_thresh = 1000
        self.is_large_image_0_thresh = 1000
        self.is_large_image_1_thresh = 1000
        self.is_large_image_low_0_thresh = 100
        self.is_large_image_low_1_thresh = 100
        self.width_thresh = 85
        self.height_thresh = 85
        self.color_thresh = 0
        self.adobeKeyUtils =  adobeKeyUtils()

    def create_folder_if_not_exists(self, blob_service_client, folder_name):
        container_client = blob_service_client.get_container_client(
            container_name)
        blob_client = container_client.get_blob_client(folder_name)

        try:
            blob_client.get_blob_properties()

            # Use the transaction_logger to log the successful folder existence
            transaction_logger.debug(f"Folder already exists: {folder_name}", folder_name=folder_name)
        except Exception as e:
            container_client.upload_blob(name=folder_name, data="")

            # Use the transaction_logger to log the creation of the folder
            transaction_logger.debug(f"Folder created: {folder_name}", folder_name=folder_name)

            # # Use the exception_logger to log any exceptions during folder creation
            # exception_logger.error("Error creating folder in blob storage",
            #                            extra={"folder_name": folder_name, "exception_message": str(e)})

    # def searchable_pdf_via_pytesseract(self, blob_path, output_pdf_save_path,tender_name):
    #     audit_logger.info(f"Entering in pytesract ==> {blob_path}, output_pdf_save_path => {output_pdf_save_path} of tender {tender_name}")

    #     start_time = time.time()

    #     images = convert_from_path(blob_path)

    #     pdf_writer = PdfWriter()

    #     for image in images:
    #         page = pytesseract.image_to_pdf_or_hocr(image, extension='pdf')
    #         pdf = PdfReader(io.BytesIO(page))
    #         pdf_writer.add_page(pdf.pages[0])

    #     # Export the searchable PDF to output_pdf_save_path
    #     with open(output_pdf_save_path, "wb") as f:
    #         pdf_writer.write(f)

    #     audit_logger.info(f'Time taken: {time.time() - start_time}')

    def save_file_to_blob_storage(self, blob_service_client, blob_path, file_data):

        df = pd.read_csv(file_data)
        container_client = blob_service_client.get_container_client(
            container_name)

        # Set the maximum number of retries
        max_retries = 20

        for attempt in range(max_retries):
            try:
                # Upload the file to blob storage
                with open(file_data, "rb") as file:
                    blob_client = container_client.get_blob_client(
                        blob=blob_path)
                    blob_client.upload_blob(file, overwrite=True)

                # Return the actual blob path
                return blob_client.url

            except Exception as e:
                # Handle the exception when saving file to blob storage fails
                exception_logger.error(f"Error saving file to blob storage (Attempt {attempt + 1}/{max_retries}): {str(e)}")
                # print(
                #     f"Error saving file to blob storage (Attempt {attempt + 1}/{max_retries}): {str(e)}")

                # Wait for 30 seconds before retrying
                time.sleep(30)

        # If all retries fail, raise the last error
        raise e

    def save_csv_data_to_mongodb(self, client, csv_path, tender_name, tender_number, published_date, tender_file_name, tender_file_path, file_type, uploaded_by, extension):
        # Connect to MongoDB
        db = client['Metadata']
        collection = db['csv_metadata']

        # Read data from CSV file
        with open(csv_path, 'r') as file:
            csv_reader = csv.DictReader(file)
            data = [dict(row, tender_name=tender_name, file_name=tender_file_name, tender_number=str(tender_number), published_date=published_date, tender_file_path=tender_file_path,
                         file_type=file_type, uploaded_by=uploaded_by, uploaded_date=datetime.now(), extension=extension, updated_by="", updated_date="") for row in csv_reader]

        # Insert data into MongoDB collection
        if data:
            collection.insert_many(data)
            transaction_logger.info("Data saved to MongoDB successfully!")
        else:
            transaction_logger.warning("No data found in the CSV file.")

    def update_blob_path_to_mongodb(self, client, tender_name, tender_number, file_name, blob_path, uploaded_by):
        
        try:
            # Connect to MongoDB
            db = client['Metadata']
            collection = db['uploaded_tender_metadata']

            # Define the filter criteria
            filter_criteria = {
                "tender_name": tender_name,
                "tender_number": tender_number,
                'tender_file_name': file_name
            }

            # Define the update operation
            update_operation = {
                "$set": {
                    "blob_csv_path": blob_path,
                    'updated_by': uploaded_by,
                    'updated_date': datetime.now()
                }
            }

            # Update the file_upload_status field for the given tender_name
            result = collection.update_one(
                filter_criteria, update_operation
            )

            if result.modified_count > 0:
                transaction_logger.info("Blob path updated successfully.")
            else:
                transaction_logger.warning(
                    "No document found to update the blob path.")
        except Exception as e:
            transaction_logger.error(f"Failed to update blob path in MongoDB: {str(e)}")
            raise e

    def download_file(self, file_url, file_name, tender_name, tender_number):
        
        folder_creation(os.path.join(self.current_path, "/mnt/supporting_folders",f"{tender_name}_{tender_number}", "splitted_pdf"))
        local_file_path = os.path.join(os.getcwd(), '/mnt/supporting_folders', f"{tender_name}_{tender_number}", 'splitted_pdf', file_name)
        blob_service_client = BlobServiceClient.from_connection_string(
            AZURE_STORAGE_CONNECTION_STRING)
        container_client = blob_service_client.get_container_client(
            container_name)
        blob_client = container_client.get_blob_client(file_url)

        attempt = 0
        while attempt < 5:
            try:
                with open(local_file_path, mode="wb") as sample_blob:
                    download_stream = blob_client.download_blob()
                    sample_blob.write(download_stream.readall())
                return local_file_path
            except Exception as e:
                transaction_logger.debug(f"Attempt {attempt + 1} failed: {e}")
                attempt += 1
                if attempt == 5:
                    raise e

        # This line is technically unnecessary, but added for clarity.
        return local_file_path



    def is_A3_or_A5(self, size, tender_name):
        audit_logger.debug(f"Checking A3 or A5 for tender_name {tender_name}")
        if size[0] > self.A3_A5_0_thresh and size[1] > self.A3_A5_1_thresh:
            return True
        else:
            return False

    def pre_process_pdfs(self, pdf_path, out_path, tender_name, tender_number, division):
        try:
            audit_logger.debug(f"=====> {pdf_path} is pre-processing for tender name {tender_name} with tender number {tender_number}.")
            
            PagesCount().extract_number_of_pages(pdf_path, division, tender_number)
            PagesCount().extract_number_of_pages_per_tender(tender_number, pdf_path)
            
            writer = PyPDF2.PdfWriter()
            doc = fitz.open(pdf_path)
            pdf = PyPDF2.PdfReader(open(pdf_path, 'rb'))
            page_list = []
            omit_list = []
            pytesseract_list = []
            normal_list = []

            for page_num in range(len(pdf.pages)):
                image_page = False
                page_fitz = doc[page_num]
                page_pypdf2 = pdf.pages[page_num]

                counter = -1
                for img_info in page_fitz.get_image_info():
                    counter += 1
                    bbox = img_info['bbox']
                    width = bbox[2] - bbox[0]
                    height = bbox[3] - bbox[1]

                    if width > self.width_thresh and height > self.height_thresh and img_info['colorspace'] > self.color_thresh:
                        audit_logger.debug(f"Page {page_num + 1} has a large image with size {(width, height)}.")
                        image_page = True

                if image_page:
                    page_list.append(page_num + 1)

            for page_num in range(len(pdf.pages)):
                image_page = False
                page_fitz = doc[page_num]
                page_pypdf2 = pdf.pages[page_num]

                if (page_num + 1) in page_list:
                    size = (page_pypdf2.mediabox.width, page_pypdf2.mediabox.height)

                    if self.is_A3_or_A5(size, tender_name):
                        audit_logger.debug(f'Removed A3/A5 page {page_num + 1} for file {pdf_path.split("/")[-1]} belonging to tender {tender_name}')
                        omit_list.append(page_num)
                        continue

                    page = doc.load_page(page_num)
                    pytesseract_list.append(page_num+1)

                    # Render the page to an image.
                    pix = page.get_pixmap()
                    data = pix.samples
                    img = Image.frombytes("RGB", [pix.width, pix.height], data)
                    page_im_tess = pytesseract.image_to_pdf_or_hocr(img, extension='pdf')
                    pdf_tess = PyPDF2.PdfReader(io.BytesIO(page_im_tess))
                    writer.add_page(pdf_tess.pages[0])
                else:
                    writer.add_page(page_pypdf2)
                    normal_list.append(page_num+1)

            if len(pdf.pages) - len(omit_list) > 0:
                with open(out_path, 'wb') as output_file:
                    writer.write(output_file)

            # TODO: pytesseract_list, normal_list can be returned for
            # future use.
            return len(pdf.pages) - len(omit_list), page_list

        except Exception as e:
            exception_logger.exception(f"Error processing PDF: {e}")
            raise  # Re-raise the exception to the calling function

    def is_readable_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    page_text = pdf_reader.pages[page_num].extract_text()
                    if page_text.strip():
                        return True
                return False
        except Exception as e:
            exception_logger.exception(f"Error checking PDF readability: {e}")
           

    def create_new_pdf_from_existing(self,input_pdf_path,output_pdf_path,tender_name,tender_number):
        
        # Open the existing PDF
        try:
            # out_pdf_file = "converted_file"+".pdf"
            # output_file_path = os.path.join(os.getcwd(), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "splitted_pdf", out_pdf_file)

            audit_logger.debug(f"Converting the pdf to readable {input_pdf_path} for tender_number {tender_number} and tender_name {tender_name}")
            with open(input_pdf_path, 'rb') as input_file:
                reader = PyPDF2.PdfReader(input_file)

                # Create a new PDF writer
                writer = PyPDF2.PdfWriter()

                # Add all pages from the existing PDF to the new PDF
                for page in reader.pages:
                    writer.add_page(page)

                # Write the content to the new PDF file
                with open(output_pdf_path, 'wb') as output_file:
                    writer.write(output_file)

            audit_logger.debug(f"New PDF file created at {output_pdf_path} for tender_number {tender_number} and tender_name {tender_name}")
            # return output_pdf_path
        except Exception as e:
            
            exception_logger.exception(f'An exception occurred at : {str(e)}')
            raise e


    def update_or_create_bad_files_dict(self,tender_number, file_name, error_message):
        try:
            # Connect to MongoDB
            client = MongoClient(MONGODB_URI)
            db = client["Metadata"]
            collection = db["tender_metadata"]

            # Find the document with the specified tender_number
            existing_doc = collection.find_one({"tender_number": tender_number})

            
            if existing_doc:
                # If a document exists, get the existing bad_file_list
                existing_bad_files = existing_doc.get("bad_file_list", [])
                audit_logger.info(f"====> 598 {existing_bad_files}")

                # Check if the file name is not already in the list before appending it
                file_exists = False
                for entry in existing_bad_files:
                    if "file_name" in entry and entry["file_name"] == file_name:
                        file_exists = True
                        break

                if not file_exists:
                    file_entry = {"file_name": file_name, "error_message": error_message}
                    existing_bad_files.append(file_entry)
                    audit_logger.info(f'+===> {existing_bad_files}')

                    # Update the bad_file_list in the existing document
                    result = collection.update_one(
                        {"tender_number": tender_number},
                        {"$set": {"bad_file_list": existing_bad_files}},
                    )
            else:
                # If no document exists, insert a new document with the bad_file_list containing the new file and error message
                new_doc = {
                    "tender_name": "wert",
                    "tender_number": tender_number,
                    "bad_file_list": [{"file_name": file_name, "error_message": error_message}],
                    # Add other fields as needed
                }
                collection.insert_one(new_doc)
                audit_logger.sucess(f"added a new document as it does not exits in entire collection")
        except PyMongoError as e:
            # Handle MongoDB errors
            exception_logger.error(f"MongoDB error: {e}")
        except Exception as e:
            # Handle other unexpected errors
            exception_logger.error(f"An unexpected error occurred: {e}")


    # def save_error_to_mongodb_and_log(self,tender_number, tender_name, error_message):
    #     try:
    #         client,_,_,_ = self.adobeKeyUtils.connect_to_mongodb()
    #         db = client["Metadata"]
    #         collection = db["tender_error_collection"]
    #         error_document = {
    #             "tender_number": tender_number,
    #             "tender_name": tender_name,
    #             "error_message": error_message
    #         }
    #         collection.insert_one(error_document)
    #     except Exception as e:
    #         logger.error(f"Error saving error to MongoDB: {str(e)}")


class deleteFileFunc():

    def __init__(self):
        self.current_path = os.getcwd()

    def delete_blob(self, blob_path):
        """
        Delete a blob from Azure Blob Storage using its blob path.

        Parameters:
            blob_path (str): The path of the blob to be deleted.

        Returns:
            bool: True if the blob was successfully deleted, False otherwise.
        """
        try:
            blob_service_client = BlobServiceClient.from_connection_string(
                AZURE_STORAGE_CONNECTION_STRING)
            container_client = blob_service_client.get_container_client(
                container_name)
            # Delete the blob
            blob_client = container_client.get_blob_client(blob_path)
            blob_client.delete_blob()

            # Return True to indicate successful deletion
            return True

        except Exception as e:
            # Handle the exception when blob deletion fails
            exception_logger.error(f"Error deleting blob '{blob_path}': {str(e)}")
            return e

    def add_category_after_deleting(self, division, tender_name, tender_number, published_date, uploaded_by, version):
        try:

            # Use the audit_logger to log the received parameters
            audit_logger.info(f"Received parameters: tender_name={tender_name}, tender_number={tender_number}, division={division}, published_date={published_date}")

            # Create a folder name based on the provided details
            folder_name = f"{tender_name}_{tender_number}_{published_date}"
            cur_datetime = datetime.now().astimezone()

            # Connect to the MongoDB server
            client = MongoClient(MONGODB_URI)

            # Access specific collections within the database
            csv_collection = client.Metadata.csv_metadata
            output_collection = client.Metadata.output_metadata
            categories_collection = client.Metadata.categories
            categories_output_collection = client.Metadata.category_output_metadata

            # Fetch tender info and categories from the respective collections
            tender_info = list(csv_collection.find(
                {"tender_number": str(tender_number), "tender_name": tender_name}))
            categories = list(categories_collection.find({}))

            df = pd.DataFrame(tender_info)

            final_document_detail = []
            file_document_detail_obj = {}

            # Initialize the Azure Blob Storage client
            blob_service_client = BlobServiceClient.from_connection_string(
                AZURE_STORAGE_CONNECTION_STRING
            )

            csvGeneratorFunc().create_folder_if_not_exists(
                blob_service_client,f"{division}/{folder_name}/output_files/parameter_wise_doc/{cur_datetime}/",
            )

            csvGeneratorFunc().create_folder_if_not_exists(
                blob_service_client, f"{division}/{folder_name}/output_files/category_wise_doc/{cur_datetime}/",
            )

            for idx, c in enumerate(categories):
                if idx == 1:
                    break

                c_final_res = {
                    "category": c["category"],
                    "pdf_file": "", "parameters": []
                }

                file_document_detail_obj[c["category"]] = {}

                # Dictionary to hold the pages and pdf names from which we need
                # to slice pages
                titles = []
                filtered_rows_other_formats = []

                for x in c["sub_category"]:

                    # Summarize output using GPT
                    summarized = category_add_func().summarized_gpt(
                        df, x["summarization_prompt"], x["gpt_prompt"]
                    )

                    # Extract the titles that are referred to for GPT output,
                    # we'll later extract all the pages under these titles
                    for z in summarized["filtered"]:
                        # if the row does not belong to a pdf file then we need
                        # to append it to category wise document directly and
                        # not slice the page.
                        if z["extension"].lower() == "pdf":
                            titles.append(z["title"])
                        else:
                            filtered_rows_other_formats.append(z)

                    # Create a document for this parameter with the data GPT returns
                    docx_file_name = (
                        x["parameter"] + "-" + cur_datetime.isoformat() +
                        ".docx"
                    )
                    folder_creation(os.path.join(
                        os.getcwd(), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file"))
                    parameter_doc = os.path.join(
                        os.getcwd(), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file", docx_file_name)
                    document = Document()
                    document.add_heading(c["category"], level=1)
                    document.add_heading(x["parameter"], level=2)
                    document.add_paragraph(cur_datetime.isoformat())
                    document.add_paragraph(summarized["gpt_output"])
                    document.save(parameter_doc)
                    parameter_blob_storage_string = f"{division}/{folder_name}/output_files/parameter_wise_doc/{cur_datetime}/{docx_file_name}"

                    docx_blob_path = category_add_func().saving_docx_to_blob(
                        blob_service_client, parameter_blob_storage_string, parameter_doc
                    )
                    # Save the output data to MongoDB
                    output_data = output_collection.insert_one(
                        {
                            "division": division,
                            "tender_number": tender_number,
                            "tender_name": tender_info[0]["tender_name"],
                            "gpt_output": summarized["gpt_output"],
                            "category": c["category"],
                            "parameter": x["parameter"],
                            "created_on": cur_datetime,
                            "created_by": uploaded_by,
                            "updated_on": "",
                            "updated_by": "",
                            "parameter_doc_blob_path": parameter_blob_storage_string,
                            "version": version,
                        }
                    )
                    # Append relevant data to the category response
                    c_final_res["parameters"].append(
                        {
                            "division": division,
                            "tender_number": tender_number,
                            "tender_name": tender_info[0]["tender_name"],
                            "gpt_output": summarized["gpt_output"],
                            "category": c["category"],
                            "parameter": x["parameter"],
                            "created_on": cur_datetime,
                            "parameter_doc_blob_path": parameter_blob_storage_string,
                        }
                    )
                    # Append relevant data to the file document detail object
                    file_document_detail_obj[c["category"]][x["parameter"]] = {
                        "id": str(output_data.inserted_id),
                        "division": division,
                        "tender_number": tender_number,
                        "tender_name": tender_info[0]["tender_name"],
                        "gpt_output": summarized["gpt_output"],
                        "category": c["category"],
                        "parameter": x["parameter"],
                        "created_on": cur_datetime,
                        "parameter_doc_blob_path": docx_blob_path,
                    }

                titles = set(titles)

                # Here we create a pdf for the entire category
                pdf_file_name = c["category"] + "-" + \
                    cur_datetime.isoformat() + ".pdf"

                # Extract all the titles in all the pdfs that have been used
                pages = {}

                # Now we need to extract all the page numbers under those titles
                for t in tender_info:
                    # also need to check if this row belongs to a pdf, if not
                    # then we cannot slice it.
                    if t['title'] in titles and t['extension'].lower() == "pdf":
                        if pages.get(t['tender_file_path']):
                            pages[t["tender_file_path"]].append(int(t["page"]))
                        else:
                            pages[t["tender_file_path"]] = [int(t["page"]), ]

                # Get only unique page numbers
                for file in list(pages.keys()):
                    pages[file] = list(set(pages[file]))
                    pages[file].sort()

                # Output file for this category
                output = PdfWriter()

                # Add all the pages to a pdf
                for file in list(pages.keys()):
                    local_file = csvGeneratorFunc().download_file(
                        file, "f-" + file.split("/")[:-1][0]
                    )
                    file_reader = PdfReader(open(local_file, "rb"))
                    for p in pages[file]:
                        output.add_page(file_reader.pages[p])

                category_pdf = os.path.join(
                    os.getcwd(
                    ), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file", c["category"] + "_pdf.pdf"
                )
                outputStream = open(category_pdf, "wb")
                output.write(outputStream)
                outputStream.close()

                # after adding sliced pages from other pdfs, we'll add other
                # data that was referred to another pdf.
                doc_data = ""
                for r in filtered_rows_other_formats:
                    doc_data += "\n\n" + r['text']

                folder_creation(os.path.join(
                    os.getcwd(), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file"))
                category_doc = os.path.join(
                    os.getcwd(
                    ), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file", c["category"] + "_docx_xlsx.pdf"
                )
                create_pdf(category_doc, doc_data)

                # Path of merged category wise pdf.
                category_final_pdf = os.path.join(
                    os.getcwd(
                    ), "/mnt/supporting_folders", f"{tender_name}_{tender_number}", "docx_file", c["category"] + ".pdf"
                )

                # Merge both pdfs.
                merger = PdfMerger()
                merger.append(category_pdf)
                merger.append(category_doc)
                merger.write(category_final_pdf)

                category_blob_storage_string = (f"{division}/{folder_name}/output_files/category_wise_doc/")
                csvGeneratorFunc().create_folder_if_not_exists(
                    blob_service_client, category_blob_storage_string
                )
                csvGeneratorFunc().create_folder_if_not_exists(
                    blob_service_client, f"{category_blob_storage_string}/{cur_datetime}/")
                # Generate the PDF file path
                category_wise_pdf_path = f"{division}/{folder_name}/output_files/category_wise_doc/{cur_datetime}/{pdf_file_name}"
                # Save the category output data to MongoDB
                cat_output = categories_output_collection.insert_one(
                    {
                        "tender_name": tender_name,
                        "category_name": c["category"],
                        "tender_number": tender_number,
                        "pdf_file_path": category_wise_pdf_path,
                        'created_on': cur_datetime,
                        'created_by': uploaded_by,
                        "updated_on": "",
                        "updated_by": "",
                    }
                )
                # Update the file_document_detail_obj with category output id
                file_document_detail_obj[c["category"]]["category_output_id"] = str(
                    cat_output.inserted_id
                )
                pdf_blob_path = category_add_func().saving_docx_to_blob(
                    blob_service_client, category_wise_pdf_path, category_final_pdf
                )
                # Update the category  with the pdf_blob_path
                c_final_res["pdf_file"] = pdf_blob_path
                # Append the final result for this category to the list
                final_document_detail.append(c_final_res)

            # Update tender metadata with the status
            uploadFileFunc().update_tender_metadata(
                tender_name, tender_number, division, uploaded_by, tender_status="Succeeded"
            )

            category_add_func().update_file_processing_status(
                tender_name, tender_number, uploaded_by, file_processing_status='Succeeded')

            # Delete all temporary files created.
            os.remove(parameter_doc)
            os.remove(category_doc)
            os.remove(category_pdf)
            os.remove(category_final_pdf)

        except Exception as e:
            # Use the exception_logger to log any exceptions
            exception_logger.error(
                "An error occurred during category wise document creation",
                extra={"exception_message": str(e)},
            )

def find_application_type(file_extension):
    mime_type, _ = mimetypes.guess_type(f"filename.{file_extension}")
    return mime_type
